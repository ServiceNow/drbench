import json
import logging
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Tuple

import numpy as np
import pandas as pd
import pymupdf as fitz
import requests
from bs4 import BeautifulSoup
from docx import Document
from openai import OpenAI

from drbench import config

# Configure logging
httpx_logger = logging.getLogger("httpx")
httpx_logger.setLevel(logging.WARNING)
logger = logging.getLogger("source_reader")


# Suppress warnings
# warnings.filterwarnings("ignore")

OPENAI_MODELS = ["gpt-4o", "gpt-4o-mini", "gpt-4.1"]

logger = logging.getLogger(__name__)


def prompt_llm(prompt, model="together_ai/meta-llama/Meta-Llama-3-8B-Instruct-Lite", **kwargs):
    show_cost = kwargs.pop("show_cost", False)
    if model in OPENAI_MODELS:
        client = OpenAI(api_key=config.OPENAI_API_KEY)
        if "response_format" in kwargs.keys():
            response = client.responses.parse(
                model=model,
                input=[{"role": "user", "content": prompt}],
                text_format=kwargs["response_format"],
            )
            return response.output_parsed
        llm = (
            lambda content: client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": content}],
                **kwargs,
            )
            .choices[0]
            .message.content
        )

        return llm(prompt)

    elif model == "vllm":
        DEFAULT_VLLM_MODEL = "neuralmagic/Meta-Llama-405B-Instruct-FP8"
        vllm_api_url = kwargs.pop("vllm_api_url", config.VLLM_API_URL or "http://localhost:8000")
        vllm_api_key = kwargs.pop("vllm_api_key", config.VLLM_API_KEY)
        vllm_model = kwargs.pop("vllm_model", config.VLLM_MODEL or DEFAULT_VLLM_MODEL)

        client = OpenAI(base_url=f"{vllm_api_url}/v1", api_key=vllm_api_key)

        messages = [
            {"role": "user", "content": prompt},
        ]

        if "response_format" in kwargs.keys():
            response = client.beta.chat.completions.parse(
                model=vllm_model,
                messages=messages,
                response_format=kwargs["response_format"],
            )
            return response.choices[0].message.parsed
        response = client.chat.completions.create(
            model=vllm_model,
            messages=messages,
            **kwargs,
        )
        return response.choices[0].message.content

    elif model.startswith("azure:"):
        from openai import AzureOpenAI

        azure_api_key = config.AZURE_API_KEY
        azure_endpoint = config.AZURE_ENDPOINT
        azure_api_version = config.AZURE_API_VERSION
        azure_deployment = model.split(":", 1)[1]  # Extract deployment name

        client = AzureOpenAI(
            api_version=azure_api_version,
            api_key=azure_api_key,
            azure_endpoint=azure_endpoint,
        )

        response = client.chat.completions.create(
            model=azure_deployment,
            messages=[{"role": "user", "content": prompt}],
            **kwargs,
        )
        return response.choices[0].message.content

    elif model.startswith("openrouter/"):
        openrouter_api_url = kwargs.pop("openrouter_api_url", config.OPENROUTER_API_URL)
        openrouter_api_key = kwargs.pop("openrouter_api_key", config.OPENROUTER_API_KEY)
        openrouter_model = model.removeprefix("openrouter/").strip()

        client = OpenAI(base_url=f"{openrouter_api_url}/v1", api_key=openrouter_api_key)

        client = OpenAI(
            base_url=openrouter_api_url,
            api_key=openrouter_api_key,
        )

        messages = [
            {"role": "user", "content": prompt},
        ]

        if "response_format" in kwargs.keys():
            response = client.beta.chat.completions.parse(
                model=openrouter_model,
                messages=messages,
                response_format=kwargs["response_format"],
            )
            return response.choices[0].message.parsed
        response = client.chat.completions.create(
            model=openrouter_model,
            messages=messages,
            **kwargs,
        )
        return response.choices[0].message.content

    else:
        import litellm
        from litellm import completion

        litellm._logging.handler.setLevel(logger.getEffectiveLevel())

        # Calculate the number of tokens
        tokens = len(prompt.split())

        # Calculate and print estimated cost for each model
        if show_cost:
            logger.debug(f"\nNumber of tokens: {tokens}")
            cost = (0.1 / 1_000_000) * tokens
            logger.debug(f"Estimated cost for {model}: ${cost:.10f}\n")

        # Make the API call using LiteLLM
        response = completion(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            **kwargs,
        )
        return response.choices[0].message.content


def break_report_to_insights(
    report_text,
    model="together_ai/meta-llama/Meta-Llama-3-8B-Instruct-Lite",
    save_json=False,
    max_retries=3,
    **kwargs,
):
    insight_claims_prompt = f"""
    Please break down the following report text into insight claims. Each insight claim should be:
    1. A single insight, that might include multiple statements and claims
    2. Independent and self-contained
    3. Each claim can have more than one sentence, but should be focused on a single insight
    4. Support each insight with citations from the report text following these specific rules:
       - Usually at the end of the report there is a list of citations with numbers
       - In the report text, citations are referenced with numbers in square brackets like [1], [2], [^1], etc.
       - When providing citations, write the actual name of the cited documents (file names or URLs), NOT the reference numbers
       - Do NOT find citations based on your general knowledge - only use citations that are explicitly presented in the report
       - For each insight you extract, look for citation markers in that specific text section and find the corresponding citation names from the reference list
       - If no citations are found for a specific insight, leave citations empty
    5. Citations should be in one of these formats (various formats will be automatically normalized):
       - Valid URLs: "https://www.example.com/article", "https://techcrunch.com/2023/report.html"
       - File names: "quarterly_report.pdf", "market_analysis.docx", "shared/file.pdf"
       - MatterMost chats: "MatterMost-Channel-Team-User" or natural descriptions like "Mattermost Message - Enterprise Chat (User: john.doe, Team: Compliance, Channel: General)"
       - Email messages: "RoundCube-from@email-to@emails-Subject" or natural descriptions like "Email from sarah.johnson@company.com on 20 Jan 2025"
       - NOTE: Various citation formats are supported and will be automatically normalized during evaluation
    6. Do not include general summaries, opinions, or claims that lack citation, just the sentences that are facts.
    7. Each claim should be a concise but complete sentence.

    ## Report text:
    <START OF REPORT>
    {report_text}
    <END OF REPORT>

    ## Output format:
    Please return the insight claims as a JSON array. For example:
    [
        {{
            "claim": "The company's revenue increased by 15% in Q3 2023",
            "citations": ["https://techfundingnews.com/salesforce-acquires-convergence-how-it-competes-with-openai-anthropic-and-googles-duet-ai/"]
        }},
        {{
            "claim": "The new product launch contributed to the growth",
            "citations": ["product_launch.pdf", "quarterly_report.pdf"]
        }},
        {{
            "claim": "Compliance team discussed FSMA requirements",
            "citations": ["MatterMost-fsma_compliance-compliance_team-john.doe"]
        }},
        {{
            "claim": "Budget review meeting was scheduled for next quarter",
            "citations": ["RoundCube-jason.kim@leesmarket.com-emily.patel@leesmarket.com,sophia.lee@leesmarket.com-Q2 Budget Review and Planning"]
        }},
        {{
            "claim": "This claim has no supporting citations",
            "citations": []
        }}
    ]

    Return only valid JSON, no additional text.

    ### Just use the report between <START OF REPORT> and <END OF REPORT> tags to generate insights.
    ### If no insights found, return an empty JSON array: []
    ### Do not use the example outputs as report content.
    """  # noqa: E501

    for attempt in range(max_retries):
        try:
            insight_claims_reponse = prompt_llm(insight_claims_prompt, model, temperature=0, **kwargs)

            # Parse the JSON response
            try:
                # Clean the response - remove any text before or after JSON
                response_clean = insight_claims_reponse.strip()

                # Find JSON array in response (handle cases where model adds extra text)
                start_idx = response_clean.find("[")
                end_idx = response_clean.rfind("]")

                if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                    json_str = response_clean[start_idx : end_idx + 1]
                    insight_list = json.loads(json_str)

                    # Validate that it's a list
                    if not isinstance(insight_list, list):
                        raise ValueError("Response is not a JSON array")

                    # Validate structure of each insight
                    validated_insights = []
                    for insight in insight_list:
                        if isinstance(insight, dict) and "claim" in insight and "citations" in insight:
                            claim = insight["claim"].strip() if isinstance(insight["claim"], str) else ""
                            citations = insight["citations"] if isinstance(insight["citations"], list) else []

                            # Basic validation - skip empty or very short claims
                            if len(claim) >= 10:
                                validated_insights.append({"claim": claim, "citations": citations})

                    # save into .json file
                    if save_json:
                        with open("insight_claims.json", "w") as f:
                            json.dump(validated_insights, f, indent=4)

                    return validated_insights
                else:
                    # No valid JSON array found
                    logger.warning(f"No valid JSON array found in response (attempt {attempt + 1})")

            except (json.JSONDecodeError, ValueError) as e:
                logger.warning(f"Failed to parse JSON response (attempt {attempt + 1}): {e}")
                logger.warning(f"Response was: {insight_claims_reponse}")

        except Exception as e:
            logger.warning(f"Error in insight extraction (attempt {attempt + 1}): {e}")

        # Don't modify prompt for retries as requested

    # If all retries failed, return empty list
    logger.error(f"Failed to extract insights after {max_retries} attempts")
    return []


def clean_citation(citation):
    """Clean and normalize citation to expected format."""
    if citation is None:
        return None

    # First try to normalize using the new flexible normalizer
    try:
        from drbench.agents.citation_normalizer import normalize_citation

        normalized = normalize_citation(citation)
        if normalized:
            # Now process the normalized citation through the existing logic
            return _process_normalized_citation(normalized)
    except ImportError:
        logger.warning("citation_normalizer not available, falling back to legacy parsing")
    except Exception as e:
        logger.debug(f"Error in citation normalization: {e}, falling back to legacy parsing")

    # Fallback to legacy parsing logic
    return _legacy_clean_citation(citation)


def _process_normalized_citation(citation):
    """Process normalized citation to the internal format expected by get_content."""
    if citation is None:
        return None

    # Return as-is if it's a URL
    if citation.startswith("http"):
        return citation

    # Handle MatterMost citation (case-insensitive check)
    if citation.lower().startswith("mattermost_"):
        citation_lower = citation.lower()
        parts = citation_lower.split("_")

        if len(parts) >= 4:
            # Format: mattermost-channel-team-user
            channel_name = parts[1]
            team_name = parts[2]
            user_name = "-".join(parts[3:])  # Join remaining parts for user (in case user has hyphens)
            return f"mattermost<sep>{channel_name}<sep>{team_name}<sep>{user_name}"
        else:
            logger.warning(f"⚠️  WARNING: Normalized MatterMost citation has insufficient parts: {citation}")
            return None

    # Handle RoundCube citation (case-insensitive check)
    if citation.lower().startswith("roundcube-"):
        citation_lower = citation.lower()
        parts = citation_lower.split("-", 3)  # Split into max 4 parts to handle subject with hyphens

        if len(parts) >= 4:
            # Format: roundcube-from-to-subject
            from_email = parts[1]
            to_emails = parts[2]
            subject = parts[3]  # Subject can contain hyphens
            return f"roundcube<sep>{from_email}<sep>{to_emails}<sep>{subject}"
        else:
            logger.warning(f"⚠️  WARNING: Normalized RoundCube citation has insufficient parts: {citation}")
            return None

    # For file citations, return as-is (already normalized to filename)
    return citation


def _legacy_clean_citation(citation):
    """Legacy citation cleaning logic (fallback)."""
    if citation is None:
        return None

    # Return as-is if it's a URL
    if citation.startswith("http"):
        return citation

    # Handle MatterMost citation (case-insensitive check)
    if citation.lower().startswith("mattermost-"):
        citation_lower = citation.lower()
        parts = citation_lower.split("-")

        if len(parts) > 4:
            # Print warning and use LLM to extract information
            logger.warning(f"MatterMost citation has more than 4 parts, using LLM to parse: {citation}")

            llm_prompt = f"""
            Parse the following MatterMost citation and extract the channel name, team name, and user name.
            Citation: {citation}

            Please return only the three values separated by commas in this exact format:
            channel_name,team_name,user_name

            Do not include any other text or formatting.
            """

            try:
                response = prompt_llm(llm_prompt, model="gpt-4o-mini", temperature=0)
                parsed_parts = [part.strip() for part in response.split(",")]

                if len(parsed_parts) == 3:
                    channel_name, team_name, user_name = parsed_parts
                    return f"mattermost<sep>{channel_name}<sep>{team_name}<sep>{user_name}"
                else:
                    logger.warning(f"LLM failed to parse MatterMost citation properly: {response}")
                    return f"mattermost<sep>{parts[1]}<sep>{parts[2]}<sep>{parts[3]}"
            except Exception as e:
                logger.warning(f"Error using LLM to parse MatterMost citation: {e}")
                return f"mattermost<sep>{parts[1]}<sep>{parts[2]}<sep>{parts[3]}"
        else:
            # Standard format: mattermost-channel-team-user
            if len(parts) == 4:
                return f"mattermost<sep>{parts[1]}<sep>{parts[2]}<sep>{parts[3]}"
            else:
                logger.warning(f"⚠️  WARNING: Citation {citation} is not in the expected format for mattermost chat.")
                return None

    # Handle RoundCube citation (case-insensitive check)
    if citation.lower().startswith("roundcube-"):
        citation_lower = citation.lower()
        parts = citation_lower.split("-")

        if len(parts) > 4:
            # Keep first 3 elements and concatenate the rest
            from_email = parts[1]
            to_emails = parts[2]
            subject = "-".join(parts[3:])  # Concatenate remaining parts
            return f"roundcube<sep>{from_email}<sep>{to_emails}<sep>{subject}"
        else:
            # Standard format: roundcube-from-to-subject
            if len(parts) == 4:
                return f"roundcube<sep>{parts[1]}<sep>{parts[2]}<sep>{parts[3]}"
            else:
                logger.warning(f"⚠️  WARNING: Citation {citation} is not in the expected format for email.")
                return None

    # For file citations, extract just the filename with extension
    file_match = re.search(r"([^/\\:]+\.[^/\\:]+)", citation, re.IGNORECASE)
    if file_match is None:
        return None
    return file_match.group(1).strip()


def get_embeddings(texts, embedding_model="text-embedding-3-small", method="openai"):
    if method == "openai":
        client = OpenAI(api_key=config.OPENAI_API_KEY)
        response = client.embeddings.create(input=texts, model=embedding_model)
        embeddings = np.array([item.embedding for item in response.data])
        embeddings = np.ascontiguousarray(embeddings, dtype=np.float32)
    elif method == "sentence-transformers":
        from sentence_transformers import SentenceTransformer

        embedding_model = SentenceTransformer(embedding_model)
        embeddings = embedding_model.encode(texts)
    else:
        raise ValueError(f"Invalid embedding method: {method}")
    return embeddings


def get_most_relevant_chunks(query, content, top_k=5, chunk_size=2048, max_chunks=200, verbose=False):
    """
    Get the most relevant chunks from content using RAG retrieval.

    Args:
        query: The query/insight to search for
        content: The content to search in
        top_k: Number of top relevant chunks to return
        max_chunks: Maximum number of chunks to process (to prevent API limits)

    Returns:
        List of most relevant chunks
    """
    # Split the content into chunks of chunk_size
    chunks = [content[i : i + chunk_size] for i in range(0, len(content), chunk_size)]

    # Filter out empty chunks
    valid_chunks = [chunk.strip() for chunk in chunks if chunk.strip()]

    if not valid_chunks:
        return []

    # Limit chunks to prevent API token limits with smart sampling
    if len(valid_chunks) > max_chunks:
        if verbose:
            logger.info(f"Too many chunks ({len(valid_chunks)}), sampling {max_chunks} chunks across the document")

        # Smart sampling: take chunks distributed across the entire document
        # This ensures we get content from beginning, middle, and end
        step = len(valid_chunks) / max_chunks
        sampled_indices = [int(i * step) for i in range(max_chunks)]
        valid_chunks = [valid_chunks[i] for i in sampled_indices]

    if len(valid_chunks) <= top_k:
        return valid_chunks

    # Initialize sentence transformer model for creating embeddings
    import faiss

    # embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
    # # Create embeddings for all chunks

    chunk_embeddings = get_embeddings(valid_chunks, embedding_model="text-embedding-3-small", method="openai")

    # Set up FAISS index for similarity search
    dimension = chunk_embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)

    # Normalize embeddings for cosine similarity
    faiss.normalize_L2(chunk_embeddings)
    index.add(chunk_embeddings)

    # Create query embedding
    query_embedding = get_embeddings([query], embedding_model="text-embedding-3-small", method="openai")
    faiss.normalize_L2(query_embedding)

    # Get top similar chunks
    scores, indices = index.search(query_embedding, min(top_k, len(valid_chunks)))

    # Extract the most relevant chunks
    relevant_chunks = []
    for score, idx in zip(scores[0], indices[0]):
        if idx < len(valid_chunks):
            relevant_chunks.append(valid_chunks[idx])

    return relevant_chunks


def get_factuality_verdict_multi(insight, citations, file_list=None, model="gpt-4o-mini", max_retries=3):
    """
    Check if the insight can be answered with multiple citations/sources.

    Args:
        insight: The claim to verify
        citations: List of citations (can contain URLs or file names)
        file_list: Dictionary mapping file names to file paths
        model: LLM model to use for verification
        max_retries: Number of retry attempts

    Returns:
        Dict with is_factual, explanation, and source_details
    """
    if not citations or len(citations) == 0:
        return {
            "is_factual": False,
            "explanation": "No citations provided to verify the claim",
            "source_details": [],
        }

    # Collect content from all valid citations
    combined_content = ""
    source_details = []

    for citation in citations:
        # Clean the citation first
        cleaned_citation = clean_citation(citation)

        # If clean_citation returns "Unknown" or None, mark as unavailable
        if cleaned_citation is None:
            source_details.append({"citation": citation, "content_available": False, "content_length": 0})
            continue

        content = get_content(cleaned_citation, file_list)
        if content:
            combined_content += f"\n\nSource: {citation}\n{content}\n"
            source_details.append(
                {
                    "citation": citation,
                    "content_available": True,
                    "content_length": len(content),
                }
            )
        else:
            source_details.append({"citation": citation, "content_available": False, "content_length": 0})

    # If no content could be retrieved from any source
    if not combined_content.strip():
        return {
            "is_factual": False,
            "explanation": "No valid content could be retrieved from any of the provided citations",
            "source_details": source_details,
        }

    # Get the most relevant chunks from combined content
    relevant_chunks = get_most_relevant_chunks(insight, combined_content)
    context = "\n".join(relevant_chunks)

    # Check if any source supports the insight
    factuality_prompt = f"""
    Given the following relevant source context from multiple sources and an insight, determine if the insight is factually supported by the sources.

    Relevant Source Materials (from multiple sources):
    {context}

    Atomic Claim: {insight}

    EVALUATION CRITERIA:
    The claim is factual if the core factual content is supported by the sources. You should be strict about important details but flexible about exact wording:

    REQUIRED for TRUE:
    1. All key factual details (numbers, dates, names, percentages, specific facts) must be present in at least one source
    2. The main substance and meaning of the claim must be supported by the source contexts
    3. No part of the claim should contradict the information in any of the sources

    ACCEPTABLE variations:
    - Different wording or phrasing that conveys the same meaning
    - Paraphrasing or summarization of the source information
    - Minor linguistic differences that don't change the factual content

    Mark as FALSE if:
    - Important factual details are missing, incorrect, or unsupported across all sources
    - The claim contradicts information in any of the sources
    - The core meaning cannot be verified from any of the source contexts

    EXAMPLES:

    FACTUAL CLAIM (TRUE):
    Source A: "Company ABC reported $50M revenue in Q3 2023"
    Source B: "ABC saw 15% growth this quarter compared to Q2"
    Claim: "ABC's revenue grew by 15% to $50M in Q3 2023"
    → TRUE (facts from both sources combined)

    UNFACTUAL CLAIM (FALSE):
    Source A: "Company ABC reported $50M revenue in Q3 2023"  
    Source B: "ABC saw 10% decline this quarter"
    Claim: "ABC's revenue grew by 15% in Q3 2023"
    → FALSE (contradicts Source B)

    UNFACTUAL CLAIM (FALSE):
    Source A: "The CEO mentioned potential expansion plans"
    Claim: "The CEO confirmed definite expansion of 25% next year"
    → FALSE (adds unsupported specifics not in any source)

    Focus on the substantive factual accuracy rather than exact word-for-word matching.

    You MUST respond with either true or false under the <factual> tag.
    Then provide a brief explanation under the <explanation> tag explaining which parts are supported or not supported and from which sources.

    Format your response EXACTLY as:

    <factual>false</factual>
    <explanation>The claim is mentioned in source X... but source Y contradicts...</explanation>

    """  # noqa: E501

    for attempt in range(max_retries):
        try:
            response = prompt_llm(factuality_prompt, model, temperature=0)

            # Parse response as {factual: bool, explanation: str} with regex
            factuality_response = re.search(
                r"<factual>(.*?)</factual>\s*<explanation>(.*?)</explanation>",
                response,
                re.DOTALL,
            )

            if factuality_response:
                is_factual_text = factuality_response.group(1).strip().lower()
                explanation_text = factuality_response.group(2).strip()

                # Validate factual value
                if is_factual_text in ["true", "false"]:
                    fact_dict = {
                        "is_factual": is_factual_text == "true",
                        "explanation": explanation_text,
                        "source_details": source_details,
                    }
                    return fact_dict
                else:
                    logger.warning(f"Invalid factual value: {is_factual_text}. Retrying...")
            else:
                logger.warning(f"Failed to parse LLM response (attempt {attempt + 1}): {response}")

                # Try alternative parsing patterns
                factual_only = re.search(r"<factual>(.*?)</factual>", response, re.DOTALL)
                if factual_only:
                    is_factual_text = factual_only.group(1).strip().lower()
                    if is_factual_text in ["true", "false"]:
                        return {
                            "is_factual": is_factual_text == "true",
                            "explanation": "No detailed explanation provided by model",
                            "source_details": source_details,
                        }

                # Check for plain text true/false
                if "true" in response.lower() and "false" not in response.lower():
                    return {
                        "is_factual": True,
                        "explanation": response.strip(),
                        "source_details": source_details,
                    }
                elif "false" in response.lower() and "true" not in response.lower():
                    return {
                        "is_factual": False,
                        "explanation": response.strip(),
                        "source_details": source_details,
                    }

        except Exception as e:
            logger.warning(f"Error in multi-source factuality check (attempt {attempt + 1}): {e}")

        # If this is not the last attempt, modify prompt to be more explicit
        if attempt < max_retries - 1:
            factuality_prompt = f"""
            Given the following relevant source context from multiple sources and an insight, determine if the insight is factually supported by the sources.

            Relevant Source Materials (from multiple sources):
            {context}

            Atomic Claim: {insight}

            EVALUATION CRITERIA:
            The claim is factual if the core factual content is supported by the sources. Focus on important details rather than exact wording:

            - All key factual details (numbers, dates, names, percentages) must be present in at least one source
            - The main substance and meaning must be supported by the source contexts
            - Different wording or paraphrasing is acceptable if the meaning is preserved
            - Mark as FALSE if important details are missing, incorrect, or contradicted by any source

            EXAMPLES:
            TRUE: Source A "$50M revenue", Source B "15% growth" → Claim "revenue grew 15% to $50M"
            FALSE: Source A "$50M revenue", Source B "10% decline" → Claim "revenue grew 15%"

            IMPORTANT: You must respond in the exact format below. Do not add any other text.

            <factual>true</factual>
            <explanation>Your detailed explanation here mentioning which sources support or contradict the claim</explanation>

            Replace 'true' with 'false' if key factual details are not supported or contradicted by the sources.
            """  # noqa: E501

    # If all retries failed, return a default response
    logger.error(f"Failed to get multi-source factuality verdict after {max_retries} attempts")
    return {
        "is_factual": False,
        "explanation": "Failed to determine factuality due to parsing errors",
        "source_details": source_details,
    }


class SourceReader:
    def __init__(self):
        # Map of file extensions to parser methods
        self.parsers = {
            ".txt": self.parse_text,
            ".json": self.parse_json,
            ".jsonl": self.parse_jsonl,
            ".csv": self.parse_csv,
            ".pdf": self.parse_pdf,
            ".docx": self.parse_docx,
            ".xlsx": self.parse_xlsx,
            ".pptx": self.parse_pptx,
            ".md": self.parse_markdown,
        }

    def parse_website(self, url: str) -> Optional[Tuple[str, str]]:
        """
        Parse a website URL and extract the title and main content.
        Can handle both HTML pages and direct file links.

        Args:
            url: URL of the website or file to parse

        Returns:
            Tuple of (title, content) or None if parsing fails
        """

        # Realistic headers to avoid 403 Forbidden
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/122.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "en-US,en;q=0.9",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Referer": "https://www.google.com/",
        }

        # Use a session to persist headers/cookies
        session = requests.Session()
        session.headers.update(headers)

        # Fetch the page
        response = session.get(url)

        # Check for success
        if response.ok:
            soup = BeautifulSoup(response.text, "html.parser")

            # # Function to count sentences
            # def count_sentences(text):
            #     return len(re.findall(r"[.!?]", text))

            content_snippets = []
            for tag in ["article", "main", "body", "div", "p", "section"]:
                content_tag = soup.find(tag)
                if content_tag:
                    content = content_tag.get_text(separator=" ", strip=True)
                    if content:
                        content_snippets.append(content)

            if len(content_snippets) > 0:
                return "\n\n".join(content_snippets)

            content = soup.get_text(separator=" ", strip=True)
            if not content:
                return None

        else:
            content = None

        return content

    # # First, make a HEAD request to check content type
    # head_response = requests.head(url, timeout=10, allow_redirects=True)
    # content_type = head_response.headers.get("content-type", "").lower()

    # # Map content types to file extensions
    # content_type_map = {
    #     "application/pdf": ".pdf",
    #     "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    #     "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    #     "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
    #     "text/plain": ".txt",
    #     "text/markdown": ".md",
    #     "application/json": ".json",
    #     "text/csv": ".csv",
    #     "application/vnd.ms-excel": ".xlsx",
    # }

    # # Check if the content type corresponds to a file we can parse
    # file_extension = None
    # for ct, ext in content_type_map.items():
    #     if ct in content_type:
    #         file_extension = ext
    #         break

    # # If it's a file we can parse, download and process it
    # if file_extension and file_extension in self.parsers:
    #     try:
    #         # Download the file
    #         response = requests.get(url, timeout=30)
    #         if response.status_code != 200:
    #             logger.warning(
    #                 f"Failed to download file from URL {url} (status: {response.status_code})"
    #             )
    #             return None

    #         # Save to temporary file
    #         with tempfile.NamedTemporaryFile(
    #             suffix=file_extension, delete=False
    #         ) as temp_file:
    #             temp_file.write(response.content)
    #             temp_file_path = Path(temp_file.name)

    #         try:
    #             # Parse the downloaded file using the appropriate parser
    #             result = self.parsers[file_extension](temp_file_path)
    #             if result:
    #                 title, content = result
    #                 # Modify title to indicate it's from a URL
    #                 title = f"{title} (from {url})"
    #                 return title, content
    #         finally:
    #             # Clean up temporary file
    #             if temp_file_path.exists():
    #                 temp_file_path.unlink()

    #     except Exception as e:
    #         logger.warning(f"Error downloading or parsing file from URL {url}: {e}")
    #         # Fall through to HTML parsing as backup

    # # If not a file or file parsing failed, treat as HTML webpage
    # response = requests.get(url, timeout=10)
    # if response.status_code != 200:
    #     logger.warning(
    #         f"Failed to fetch URL {url} (status: {response.status_code})"
    #     )
    #     return None

    # soup = BeautifulSoup(response.text, "html.parser")

    # # Extract the title
    # title_tag = soup.find("title")
    # title = title_tag.get_text(strip=True) if title_tag else "No Title"

    # # Extract the main textual content from common containers
    # for tag in ["article", "main", "body", "div"]:
    #     content_tag = soup.find(tag)
    #     if content_tag:
    #         content = content_tag.get_text(separator=" ", strip=True)
    #         if content:
    #             return title, content

    # # Fallback if no tag yields content
    # content = soup.get_text(separator=" ", strip=True)
    # if not content:
    #     logger.warning(f"No content found at URL: {url}")
    #     return None

    # return content

    def parse_file(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """
        Parse a file and extract title and content.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (title, content) or None if parsing fails
        """
        try:
            # Check if file exists and is readable
            if not file_path.exists() or not file_path.is_file():
                logger.warning(f"File does not exist or is not a regular file: {file_path}")
                return None

            # Get file extension
            ext = file_path.suffix.lower()

            # Check if we have a specific parser for this extension
            if ext in self.parsers:
                return self.parsers[ext](file_path)

            logger.warning(f"Unsupported file type: {ext}")
            return None

        except Exception as e:
            logger.warning(f"Error parsing file {file_path}: {e}")
            return None

    def _describe_dataframe_to_sentences(self, df: pd.DataFrame) -> str:
        """Convert DataFrame summary statistics into human-readable sentences."""
        df_num = df.apply(lambda x: pd.to_numeric(x, errors="coerce"))
        summary_df = df_num.describe()
        lines = []
        for col in summary_df.columns:
            stats = summary_df[col]
            lines.append(f"For column '{col}':")
            lines.append(f"  - Count: {stats['count']:.0f}")
            lines.append(f"  - Mean: {stats['mean']:.2f}")
            lines.append(f"  - Standard deviation: {stats['std']:.2f}")
            lines.append(f"  - Min: {stats['min']:.2f}")
            lines.append(f"  - 25th percentile: {stats['25%']:.2f}")
            lines.append(f"  - Median (50%): {stats['50%']:.2f}")
            lines.append(f"  - 75th percentile: {stats['75%']:.2f}")
            lines.append(f"  - Max: {stats['max']:.2f}")
            lines.append("")  # Add a blank line between columns
        return "\n".join(lines)

    def parse_pdf(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """
        Parse PDF files using PyMuPDF (fitz).

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (title, content) or None if parsing fails
        """
        try:
            if not file_path.exists() or not file_path.is_file():
                logger.warning(f"File does not exist or is not a regular file: {file_path}")
                return None

            doc = fitz.open(str(file_path))
            content = ""

            for page in doc:
                content += page.get_text()

            doc.close()

            if not content.strip():
                logger.warning(f"No extractable text found in: {file_path}")
                return None

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Clean up content
            content = re.sub(r"\n{3,}", "\n\n", content.strip())

            return title, content

        except Exception as e:
            logger.warning(f"Error parsing PDF file {file_path}: {e}")
            return None

    def parse_docx(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse DOCX files using python-docx."""
        try:
            doc = Document(file_path)

            # Extract title from document properties or first heading
            title = file_path.stem.replace("_", " ").title()
            if doc.core_properties.title:
                title = doc.core_properties.title

            # Extract text from paragraphs and tables
            full_text = []

            # Get text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            # Join all text
            content = "\n\n".join(full_text)

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing DOCX file {file_path}: {e}")
            return None

    def parse_xlsx(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse XLSX files using pandas."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets = excel_file.sheet_names

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Create a detailed summary of the Excel content
            content_parts = []
            for sheet_name in sheets:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                num_rows, num_cols = df.shape

                # Add sheet summary
                content_parts.append(f"Sheet '{sheet_name}': {num_rows} rows × {num_cols} columns")

                # Add DataFrame summary statistics
                if num_rows > 0 and num_cols > 0:
                    content_parts.append("Data summary:\n")
                    content_parts.append(self._describe_dataframe_to_sentences(df))

                # Add column names
                if num_cols > 0:
                    cols = df.columns.tolist()
                    content_parts.append(f"Columns: {', '.join(str(col) for col in cols)}")

                # Add first few rows of data
                if num_rows > 0:
                    content_parts.append("\nFirst 20 rows:")
                    content_parts.append(df.head(20).to_string())

            content = "\n\n".join(content_parts)

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing XLSX file {file_path}: {e}")
            return None

    def parse_pptx(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse PPTX files by extracting text from slides."""
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                # Extract PPTX (it's a ZIP file)
                subprocess.run(["unzip", "-q", str(file_path), "-d", tmp_dir], check=True)

                # Look for slide content XML files
                slide_files = list(Path(tmp_dir).glob("ppt/slides/slide*.xml"))
                slide_texts = []

                for slide_file in sorted(slide_files):
                    with open(slide_file, "r", encoding="utf-8", errors="ignore") as f:
                        slide_xml = f.read()
                        # Extract text between <a:t> tags (text elements in PPTX XML)
                        texts = re.findall(r"<a:t>([^<]+)</a:t>", slide_xml)
                        if texts:
                            slide_texts.append(" ".join(texts))

                # Use filename as title
                title = file_path.stem.replace("_", " ").title()

                # Create content with all slides
                if slide_texts:
                    content = f"Presentation with {len(slide_texts)} slides:\n\n"
                    for i, text in enumerate(slide_texts, 1):
                        content += f"Slide {i}:\n{text}\n\n"
                else:
                    content = "Empty presentation"

                return title, content
        except Exception as e:
            logger.warning(f"Error parsing PPTX file {file_path}: {e}")
            return None

    def parse_text(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse plain text files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Clean up content
            content = re.sub(r"\n{3,}", "\n\n", content.strip())  # Replace multiple newlines with double newlines

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing text file {file_path}: {e}")
            return None

    def parse_csv(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse CSV files."""
        try:
            # Read CSV with pandas for better handling of different formats
            df = pd.read_csv(file_path)

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Create detailed content
            content_parts = []

            # Add summary
            num_rows, num_cols = df.shape
            content_parts.append(f"CSV data with {num_rows} rows and {num_cols} columns")

            # Add DataFrame summary statistics
            if num_rows > 0 and num_cols > 0:
                content_parts.append("Data summary:\n")
                content_parts.append(self._describe_dataframe_to_sentences(df))

            # Add column names
            if num_cols > 0:
                cols = df.columns.tolist()
                content_parts.append(f"Headers: {', '.join(str(col) for col in cols)}")

            # Add first few rows of data
            if num_rows > 0:
                content_parts.append("\nFirst few rows:")
                content_parts.append(df.head(5).to_string())

            content = "\n\n".join(content_parts)

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing CSV file {file_path}: {e}")
            return None

    def parse_json(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse JSON files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Try to extract meaningful content from JSON
            content = ""

            # Look for common fields that might contain descriptive text
            text_fields = [
                "description",
                "summary",
                "text",
                "content",
                "body",
                "abstract",
            ]

            if isinstance(data, dict):
                # Try to find a title field
                for title_field in ["title", "name", "heading"]:
                    if title_field in data and isinstance(data[title_field], str) and data[title_field].strip():
                        title = data[title_field].strip()
                        break

                # Try to find a content field
                for field in text_fields:
                    if field in data and isinstance(data[field], str) and data[field].strip():
                        content = data[field].strip()
                        break

            # If no content field found, use a formatted string representation of the data
            if not content:
                content = json.dumps(data, indent=2)

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing JSON file {file_path}: {e}")
            return None

    def parse_jsonl(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse JSONL files."""
        key_fields = ["direct_post", "post"]
        content = ""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()
                for line in lines:
                    if set(json.loads(line).keys()) & set(key_fields):
                        content += line

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing JSONL file {file_path}: {e}")
            return None

    def parse_markdown(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse markdown files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            # Extract title from first heading or use filename
            title = file_path.stem.replace("_", " ").title()
            lines = content.strip().split("\n")
            if lines and lines[0].startswith("# "):
                title = lines[0].lstrip("# ").strip()

            # Clean up content
            # Remove markdown headers
            content = re.sub(r"^#.*$", "", content, flags=re.MULTILINE)
            # Remove markdown links but keep the text
            content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)
            # Remove markdown emphasis
            content = re.sub(r"[*_]{1,2}([^*_]+)[*_]{1,2}", r"\1", content)
            # Remove markdown code blocks
            content = re.sub(r"```.*?```", "", content, flags=re.DOTALL)
            # Remove markdown inline code
            content = re.sub(r"`([^`]+)`", r"\1", content)
            # Remove markdown blockquotes
            content = re.sub(r"^>.*$", "", content, flags=re.MULTILINE)
            # Remove markdown horizontal rules
            content = re.sub(r"^[-*_]{3,}$", "", content, flags=re.MULTILINE)
            # Remove markdown lists
            content = re.sub(r"^[-*+]\s+", "", content, flags=re.MULTILINE)
            content = re.sub(r"^\d+\.\s+", "", content, flags=re.MULTILINE)

            # Clean up whitespace
            content = re.sub(r"\n{3,}", "\n\n", content.strip())

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing markdown file {file_path}: {e}")
            return None


def get_content(source, file_list=None):
    if str(source).lower() == "none":
        return None

    if source.startswith("http"):
        result = SourceReader().parse_website(source)
        return result if result is not None else None
    elif source.lower().startswith("mattermost"):
        # Handle MatterMost chat content retrieval
        if file_list is None:
            raise ValueError("file_list is required for MatterMost citations")

        # Parse the citation format: mattermost<sep>channel<sep>team<sep>user
        parts = source.split("<sep>")
        if len(parts) != 4:
            logger.warning(f"Invalid MatterMost citation format: {source}")
            return None

        _, channel_name, team_name, user_name = [part.lower() for part in parts]

        # Find all mattermost_chat jsonl files
        mattermost_files = [
            filename
            for filename in file_list.keys()
            if filename.lower().startswith("mattermost_chat") and filename.endswith(".jsonl")
        ]

        collected_posts = []

        for filename in mattermost_files:
            file_path = file_list[filename]
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            data = json.loads(line)
                            if (
                                data.get("type") == "post"
                                and data.get("post", {}).get("team", "").lower() == team_name
                                and data.get("post", {}).get("channel", "").lower() == channel_name
                                and data.get("post", {}).get("user", "").lower() == user_name
                            ):
                                collected_posts.append(data)
            except Exception as e:
                logger.warning(f"Error reading MatterMost file {filename}: {e}")
                continue

        if not collected_posts:
            logger.warning(f"No matching MatterMost posts found for {source}")
            return None

        # Convert collected posts to string
        content_parts = []
        for post_data in collected_posts:
            post = post_data.get("post", {})
            message = post.get("message", "")
            created_at = post.get("created_at", "")
            content_parts.append(f"Message: {message} (Created: {created_at})")

        return "\n\n".join(content_parts)

    elif source.lower().startswith("roundcube"):
        # Handle RoundCube email content retrieval
        if file_list is None:
            raise ValueError("file_list is required for RoundCube citations")

        # Parse the citation format: roundcube<sep>from_email<sep>to_emails<sep>subject
        parts = source.split("<sep>")
        if len(parts) != 4:
            logger.warning(f"Invalid RoundCube citation format: {source}")
            return None

        _, from_email, to_emails, subject = [part.lower() for part in parts]

        # Split to_emails by comma to handle multiple recipients
        to_email_list = [email.strip() for email in to_emails.split(",")]

        # Find all roundcube_email jsonl files
        roundcube_files = [
            filename
            for filename in file_list.keys()
            if filename.lower().startswith("roundcube_email") and filename.endswith(".jsonl")
        ]

        collected_emails = []

        for filename in roundcube_files:
            file_path = file_list[filename]
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            data = json.loads(line)
                            if data.get("type") == "email":
                                email_data = data
                                email_from = email_data.get("from", "").lower()
                                email_to = [email.lower() for email in email_data.get("to", [])]
                                email_subject = email_data.get("subject", "").lower()

                                # Check if from email matches
                                if email_from == from_email:
                                    # Check if subject matches
                                    if email_subject == subject:
                                        # Check if any of the to emails match
                                        if any(to_email in email_to for to_email in to_email_list):
                                            collected_emails.append(email_data)
            except Exception as e:
                logger.warning(f"Error reading RoundCube file {filename}: {e}")
                continue

        if not collected_emails:
            logger.warning(f"No matching RoundCube emails found for {source}")
            return None

        # Concatenate all email bodies
        content_parts = []
        attachments_list = []

        for email_data in collected_emails:
            body = email_data.get("body", "")
            date = email_data.get("date", "")
            from_name = email_data.get("from_name", "")
            subject_original = email_data.get("subject", "")
            attachments = email_data.get("attachments", [])

            if attachments:
                attachments_list.extend(attachments)

            email_content = f"From: {from_name} ({email_data.get('from', '')})\n"
            email_content += f"Subject: {subject_original}\n"
            email_content += f"Date: {date}\n"
            email_content += f"Body: {body}"

            content_parts.append(email_content)

        return "\n\n---EMAIL SEPARATOR---\n\n".join(content_parts)

    else:
        if file_list is None:
            raise ValueError("file_list is required")

        # extract only the string with an extension in source, any extension
        file_match = re.search(r"([^/\\:]+\.[^/\\:]+)", source, re.IGNORECASE)
        if file_match is None:
            return None

        file_name = file_match.group(1).strip()

        if file_name in file_list:
            path = file_list[file_name]
            result = SourceReader().parse_file(path)
            if result is not None:
                return result[1]  # Return content part of the tuple
            else:
                return None
        else:
            return None
