import re
from pathlib import Path
from typing import Tuple, Optional
import json
import logging
import mimetypes
import subprocess
import tempfile
import pandas as pd
from docx import Document
import requests
from bs4 import BeautifulSoup
import fitz

# Configure logging
logger = logging.getLogger("source_reader")

# Initialize mimetypes
mimetypes.init()


class SourceReader:
    def __init__(self):
        # Map of file extensions to parser methods
        self.parsers = {
            ".txt": self.parse_text,
            ".json": self.parse_json,
            ".jsonl": self.parse_jsonl,
            ".csv": self.parse_csv,
            ".pdf": self.parse_pdf,
            ".docx": self.parse_docx,
            ".xlsx": self.parse_xlsx,
            ".pptx": self.parse_pptx,
            ".md": self.parse_markdown,
        }

    def parse_website(self, url: str) -> Optional[Tuple[str, str]]:
        """
        Parse a website URL and extract the title and main content.
        Can handle both HTML pages and direct file links.

        Args:
            url: URL of the website or file to parse

        Returns:
            Tuple of (title, content) or None if parsing fails
        """

        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36"
        }

        try:
            # First, make a HEAD request to check content type
            head_response = requests.head(url, headers=headers, timeout=10, allow_redirects=True)
            content_type = head_response.headers.get("content-type", "").lower()

            # Map content types to file extensions
            content_type_map = {
                "application/pdf": ".pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
                "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
                "text/plain": ".txt",
                "text/markdown": ".md",
                "application/json": ".json",
                "text/csv": ".csv",
                "application/vnd.ms-excel": ".xlsx",
            }

            # Check if the content type corresponds to a file we can parse
            file_extension = None
            for ct, ext in content_type_map.items():
                if ct in content_type:
                    file_extension = ext
                    break

            # If it's a file we can parse, download and process it
            if file_extension and file_extension in self.parsers:
                try:
                    # Download the file
                    response = requests.get(url, headers=headers, timeout=30)
                    if response.status_code != 200:
                        logger.warning(f"Failed to download file from URL {url} (status: {response.status_code})")
                        return None

                    # Save to temporary file
                    with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = Path(temp_file.name)

                    try:
                        # Parse the downloaded file using the appropriate parser
                        result = self.parsers[file_extension](temp_file_path)
                        if result:
                            title, content = result
                            # Modify title to indicate it's from a URL
                            title = f"{title} (from {url})"
                            return title, content
                    finally:
                        # Clean up temporary file
                        if temp_file_path.exists():
                            temp_file_path.unlink()

                except Exception as e:
                    logger.warning(f"Error downloading or parsing file from URL {url}: {e}")
                    # Fall through to HTML parsing as backup

            # If not a file or file parsing failed, treat as HTML webpage
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code != 200:
                logger.warning(f"Failed to fetch URL {url} (status: {response.status_code})")
                return None

            soup = BeautifulSoup(response.text, "html.parser")

            # Extract the title
            title_tag = soup.find("title")
            title = title_tag.get_text(strip=True) if title_tag else "No Title"

            # Extract the main textual content from common containers
            for tag in ["article", "main", "body", "div"]:
                content_tag = soup.find(tag)
                if content_tag:
                    content = content_tag.get_text(separator=" ", strip=True)
                    if content:
                        return title, content

            # Fallback if no tag yields content
            content = soup.get_text(separator=" ", strip=True)
            if not content:
                logger.warning(f"No content found at URL: {url}")
                return None

            return title, content

        except Exception as e:
            logger.warning(f"Error parsing website {url}: {e}")
            return None

    def parse_file(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """
        Parse a file and extract title and content.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (title, content) or None if parsing fails
        """
        if isinstance(file_path, str):
            file_path = Path(file_path)
        try:
            # Check if file exists and is readable
            if not file_path.exists() or not file_path.is_file():
                logger.warning(f"File does not exist or is not a regular file: {file_path}")
                return None

            # Get file extension
            ext = file_path.suffix.lower()

            # Check if we have a specific parser for this extension
            if ext in self.parsers:
                return self.parsers[ext](file_path)

            logger.warning(f"Unsupported file type: {ext}")
            return None

        except Exception as e:
            logger.warning(f"Error parsing file {file_path}: {e}")
            return None

    def _describe_dataframe_to_sentences(self, df: pd.DataFrame) -> str:
        """Convert DataFrame summary statistics into human-readable sentences."""
        summary_df = df.describe()
        lines = []
        for col in summary_df.columns:
            stats = summary_df[col]
            lines.append(f"For column '{col}':")
            lines.append(f"  - Count: {stats['count']:.0f}")
            lines.append(f"  - Mean: {stats['mean']:.2f}")
            lines.append(f"  - Standard deviation: {stats['std']:.2f}")
            lines.append(f"  - Min: {stats['min']:.2f}")
            lines.append(f"  - 25th percentile: {stats['25%']:.2f}")
            lines.append(f"  - Median (50%): {stats['50%']:.2f}")
            lines.append(f"  - 75th percentile: {stats['75%']:.2f}")
            lines.append(f"  - Max: {stats['max']:.2f}")
            lines.append("")  # Add a blank line between columns
        return "\n".join(lines)

    def parse_pdf(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """
        Parse PDF files using PyMuPDF (fitz).

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (title, content) or None if parsing fails
        """
        try:
            if not file_path.exists() or not file_path.is_file():
                logger.warning(f"File does not exist or is not a regular file: {file_path}")
                return None

            doc = fitz.open(str(file_path))
            content = ""

            for page in doc:
                content += page.get_text()

            doc.close()

            if not content.strip():
                logger.warning(f"No extractable text found in: {file_path}")
                return None

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Clean up content
            content = re.sub(r"\n{3,}", "\n\n", content.strip())

            return title, content

        except Exception as e:
            logger.warning(f"Error parsing PDF file {file_path}: {e}")
            return None

    def parse_docx(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse DOCX files using python-docx."""
        try:
            doc = Document(file_path)

            # Extract title from document properties or first heading
            title = file_path.stem.replace("_", " ").title()
            if doc.core_properties.title:
                title = doc.core_properties.title

            # Extract text from paragraphs and tables
            full_text = []

            # Get text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            # Join all text
            content = "\n\n".join(full_text)

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing DOCX file {file_path}: {e}")
            return None

    def parse_xlsx(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse XLSX files using pandas."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets = excel_file.sheet_names

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Create a detailed summary of the Excel content
            content_parts = []
            for sheet_name in sheets:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                num_rows, num_cols = df.shape

                # Add sheet summary
                content_parts.append(f"Sheet '{sheet_name}': {num_rows} rows × {num_cols} columns")

                # Add DataFrame summary statistics
                if num_rows > 0 and num_cols > 0:
                    content_parts.append(f"Data summary:\n")
                    content_parts.append(self._describe_dataframe_to_sentences(df))

                # Add column names
                if num_cols > 0:
                    cols = df.columns.tolist()
                    content_parts.append(f"Columns: {', '.join(str(col) for col in cols)}")

                # Add first few rows of data
                if num_rows > 0:
                    content_parts.append("\nFirst few rows:")
                    content_parts.append(df.head(5).to_string())

            content = "\n\n".join(content_parts)

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing XLSX file {file_path}: {e}")
            return None

    def parse_pptx(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse PPTX files by extracting text from slides."""
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                # Extract PPTX (it's a ZIP file)
                subprocess.run(["unzip", "-q", str(file_path), "-d", tmp_dir], check=True)

                # Look for slide content XML files
                slide_files = list(Path(tmp_dir).glob("ppt/slides/slide*.xml"))
                slide_texts = []

                for slide_file in sorted(slide_files):
                    with open(slide_file, "r", encoding="utf-8", errors="ignore") as f:
                        slide_xml = f.read()
                        # Extract text between <a:t> tags (text elements in PPTX XML)
                        texts = re.findall(r"<a:t>([^<]+)</a:t>", slide_xml)
                        if texts:
                            slide_texts.append(" ".join(texts))

                # Use filename as title
                title = file_path.stem.replace("_", " ").title()

                # Create content with all slides
                if slide_texts:
                    content = f"Presentation with {len(slide_texts)} slides:\n\n"
                    for i, text in enumerate(slide_texts, 1):
                        content += f"Slide {i}:\n{text}\n\n"
                else:
                    content = "Empty presentation"

                return title, content
        except Exception as e:
            logger.warning(f"Error parsing PPTX file {file_path}: {e}")
            return None

    def parse_text(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse plain text files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Clean up content
            content = re.sub(r"\n{3,}", "\n\n", content.strip())  # Replace multiple newlines with double newlines

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing text file {file_path}: {e}")
            return None

    def parse_csv(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse CSV files."""
        try:
            # Read CSV with pandas for better handling of different formats
            df = pd.read_csv(file_path)

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Create detailed content
            content_parts = []

            # Add summary
            num_rows, num_cols = df.shape
            content_parts.append(f"CSV data with {num_rows} rows and {num_cols} columns")

            # Add DataFrame summary statistics
            if num_rows > 0 and num_cols > 0:
                content_parts.append(f"Data summary:\n")
                content_parts.append(self._describe_dataframe_to_sentences(df))

            # Add column names
            if num_cols > 0:
                cols = df.columns.tolist()
                content_parts.append(f"Headers: {', '.join(str(col) for col in cols)}")

            # Add first few rows of data
            if num_rows > 0:
                content_parts.append("\nFirst few rows:")
                content_parts.append(df.head(5).to_string())

            content = "\n\n".join(content_parts)

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing CSV file {file_path}: {e}")
            return None

    def parse_json(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse JSON files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            # Try to extract meaningful content from JSON
            content = ""

            # Look for common fields that might contain descriptive text
            text_fields = [
                "description",
                "summary",
                "text",
                "content",
                "body",
                "abstract",
            ]

            if isinstance(data, dict):
                # Try to find a title field
                for title_field in ["title", "name", "heading"]:
                    if title_field in data and isinstance(data[title_field], str) and data[title_field].strip():
                        title = data[title_field].strip()
                        break

                # Try to find a content field
                for field in text_fields:
                    if field in data and isinstance(data[field], str) and data[field].strip():
                        content = data[field].strip()
                        break

            # If no content field found, use a formatted string representation of the data
            if not content:
                content = json.dumps(data, indent=2)

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing JSON file {file_path}: {e}")
            return None

    def parse_jsonl(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse JSONL files."""
        key_fields = ["direct_post", "post"]
        content = ""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()
                for line in lines:
                    if set(json.loads(line).keys()) & set(key_fields):
                        content += line

            # Use filename as title
            title = file_path.stem.replace("_", " ").title()

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing JSONL file {file_path}: {e}")
            return None

    def parse_markdown(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse markdown files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            # Extract title from first heading or use filename
            title = file_path.stem.replace("_", " ").title()
            lines = content.strip().split("\n")
            if lines and lines[0].startswith("# "):
                title = lines[0].lstrip("# ").strip()

            # Clean up content
            # Remove markdown headers
            content = re.sub(r"^#.*$", "", content, flags=re.MULTILINE)
            # Remove markdown links but keep the text
            content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)
            # Remove markdown emphasis
            content = re.sub(r"[*_]{1,2}([^*_]+)[*_]{1,2}", r"\1", content)
            # Remove markdown code blocks
            content = re.sub(r"```.*?```", "", content, flags=re.DOTALL)
            # Remove markdown inline code
            content = re.sub(r"`([^`]+)`", r"\1", content)
            # Remove markdown blockquotes
            content = re.sub(r"^>.*$", "", content, flags=re.MULTILINE)
            # Remove markdown horizontal rules
            content = re.sub(r"^[-*_]{3,}$", "", content, flags=re.MULTILINE)
            # Remove markdown lists
            content = re.sub(r"^[-*+]\s+", "", content, flags=re.MULTILINE)
            content = re.sub(r"^\d+\.\s+", "", content, flags=re.MULTILINE)

            # Clean up whitespace
            content = re.sub(r"\n{3,}", "\n\n", content.strip())

            return title, content
        except Exception as e:
            logger.warning(f"Error parsing markdown file {file_path}: {e}")
            return None
